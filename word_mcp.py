import csv
import difflib
import os
import json
import re
import jieba

from collections import Counter
from mcp.server.fastmcp import FastMCP
from typing import List
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from reportlab.lib.fonts import addMapping
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.units import inch
from PyPDF2 import PdfReader

mcp = FastMCP("word_mcp", log_level="ERROR")

# 导入中文字体, 解决中文乱码问题
pdfmetrics.registerFont(TTFont('SimHei', 'simhei.ttf'))  # 黑体
addMapping('SimHei', 0, 0, 'SimHei')


@mcp.tool()
def create_empty_txt(filename: str) -> str | dict[str, str]:
    """
    在指定路径上创建一个空白的TXT文件
    """
    if not filename.lower().endswith('.txt'):
        filename += '.txt'

    output_path = os.environ.get('WORD_MCP_PATH')  # 从环境变量获取输出路径，如果未设置则使用默认桌面路径
    if not output_path:
        output_path = os.path.join(os.path.expanduser('~'), '桌面')

    file_path = os.path.join(output_path, filename)

    try:
        # 创建输出目录（如果不存在）
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # 创建空白文件
        with open(file_path, 'w', encoding='utf-8') as f:
            pass
        return response_handler({"status": "success", "message": f"成功在 {output_path} 创建了空白文件: {filename}"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"创建文件时出错: {str(e)}"})


@mcp.tool()
def create_word_document(filename: str) -> str | str | dict[str, str]:
    """
    创建一个新的Word文档
    """
    # 确保文件名有.docx扩展名
    if not filename.lower().endswith('.docx'):
        filename += '.docx'

    # 从环境变量获取输出路径，如果未设置则使用默认桌面路径
    output_path = os.environ.get('WORD_MCP_PATH')  # 保持环境变量名不变，以兼容现有配置
    if not output_path:
        output_path = os.path.join(os.path.expanduser('~'), '桌面')

    file_path = os.path.join(output_path, filename)

    try:
        # 创建输出目录（如果不存在）
        os.makedirs(os.path.dirname(file_path), exist_ok=True)

        # 创建新的Word文档
        doc = Document()

        # 保存文档
        doc.save(file_path)

        return response_handler({"status": "success", "message": f"成功创建文件: {filename}"})
    except Exception as e:
        return response_handler({"status": "error", "message": str(e)})


@mcp.tool()
def open_and_read_word_document(file_path: str) -> str:
    """
    打开并读取Word文档,返回文档信息头和内容
    """
    # 是否提供了完整路径
    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"错误: 文件 {file_path} 不存在"})

    try:
        doc = Document(file_path)

        # 提取文档基本信息
        paragraphs = [p.text for p in doc.paragraphs]
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith('Heading')]

        # 构建文档信息头
        doc_info = (
            f"文件名: {os.path.basename(file_path)}\n"
            f"段落数: {len(paragraphs)}\n"
            f"标题数: {len(headings)}\n\n"
        )

        # 构建完整文档内容，保留段落结构，并在每段前添加段落编号
        full_content = ""
        for i, p_text in enumerate(paragraphs):
            full_content += f"[{i}] {p_text}\n"

        return response_handler({"status": "success", "message": doc_info + full_content})
    except Exception as e:
        return response_handler({"status": "error", "message": f"读取Word文档时出错: {str(e)}"})


@mcp.tool()
def query_document_info(file_path: str) -> str:
    """
    查询Word文档的基本信息
    :param file_path: 文件路径
    :return: 包含以下字段的JSON:
        - 创建者, 创建时间, 最后修改者, 最后修改时间
        - 页数, 字数, 段落数, 表格数, 图片数
        - 标题, 作者, 关键字, 主题, 备注
        - 文件大小, 文件格式版本
    """

    # 路径处理
    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH', os.path.expanduser('~/桌面'))
        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": "文件不存在"})

    try:
        # 读取基础信息
        doc = Document(file_path)
        core_props = doc.core_properties
        stats = {
            'paragraphs': len(doc.paragraphs),
            'tables': len(doc.tables),
            'images': len(doc.inline_shapes)
        }

        # 获取文件系统信息
        file_stat = os.stat(file_path)
        file_size = f"{file_stat.st_size / 1024:.2f} KB"

        # 判断文件版本
        file_version = "DOCX" if file_path.lower().endswith('.docx') else "Legacy DOC"

        # 构建元数据字典
        metadata = {
            "标题": core_props.title,
            "创建者": core_props.author,
            "创建时间": core_props.created.strftime("%Y-%m-%d %H:%M:%S") if core_props.created else None,
            "最后修改者": core_props.last_modified_by,
            "最后修改时间": core_props.modified.strftime("%Y-%m-%d %H:%M:%S") if core_props.modified else None,
            "备注": core_props.comments,
            "大小": file_size,
            "文件格式": file_version
        }

        # 合并统计数据
        result = {**metadata, **stats}
        return response_handler({"status": "success", "data": result})

    except Exception as e:
        return response_handler({"status": "error", "message": f"文档解析失败: {str(e)}"})


@mcp.tool()
def add_text_to_document(
        file_path: str,
        text: str,
        is_heading: bool = False,
        heading_level: int = 1,
        alignment: str = "left",
        paragraph_index: int = -1,
        direction: str = "behind"
) -> str:
    """
    向Word文档添加文本内容
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param text: 要添加的文本内容
    :param is_heading: 是否作为标题添加
    :param heading_level: 标题级别 (1-9)，仅当is_heading=True时有效
    :param alignment: 对齐方式，可选值: "left", "center", "right", "justify"
    :param paragraph_index: 段落索引 (从0开始计数)，指定添加内容的位置，默认为-1表示文档末尾
    :param direction: 指定内容添加的位置，可选值: "front"(段落前), "behind"(段落后)，默认为behind
    :
    """

    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"错误: 文件 {file_path} 不存在"})

    if is_heading and (heading_level < 1 or heading_level > 9):
        return response_handler({"status": "error", "message": "错误: 标题级别必须在1至9之间"})

    if direction not in ["front", "behind"]:
        return response_handler({"status": "error", "message": f"错误: 无效的方向参数 '{direction}'，可选值为: front, behind"})

    # 映射对齐方式参数
    alignment_map = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
    }

    if alignment not in alignment_map:
        return response_handler(
            {"status": "error", "message": f"错误: 不支持的对齐方式 '{alignment}'，可选值为: left, center, right, justify"})

    try:
        doc = Document(file_path)

        # 检查段落索引是否有效（当不为默认值-1时）
        if paragraph_index != -1 and (paragraph_index < 0 or paragraph_index >= len(doc.paragraphs)):
            return response_handler(
                {"status": "error", "message": f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"})

        # 创建新段落或标题
        if is_heading:
            # 创建标题
            new_paragraph = doc.add_heading(text, level=heading_level)
            # 设置标题的对齐方式
            new_paragraph.alignment = alignment_map[alignment]
        else:
            # 创建普通段落
            new_paragraph = doc.add_paragraph(text)
            new_paragraph.alignment = alignment_map[alignment]

        # 如果指定了段落索引，根据direction参数调整新段落的位置
        if paragraph_index != -1:
            target_paragraph = doc.paragraphs[paragraph_index]

            # 获取新段落的XML元素
            new_p = new_paragraph._p

            # 删除刚刚添加的段落，因为我们需要重新插入到指定位置
            new_p.getparent().remove(new_p)

            if direction == "front":
                # 在目标段落前插入
                target_paragraph._p.addprevious(new_p)
            else:  # direction == "behind"
                # 在目标段落后插入
                target_paragraph._p.addnext(new_p)

        # 保存文档
        doc.save(file_path)

        # 构建返回消息
        if paragraph_index == -1:
            position_msg = "文档末尾"
        else:
            position_msg = f"第 {paragraph_index + 1} 段落{'前' if direction == 'front' else '后'}"

        return response_handler({"status": "success", "message": f"成功在{position_msg}添加了{'标题' if is_heading else '文本'}"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"向Word文档添加内容时出错: {str(e)}"})


@mcp.tool()
def format_text_in_document(
        file_path: str,
        paragraph_index: int,
        font_name: str = None,
        font_size: int = None,
        bold: bool = False,
        italic: bool = False,
        underline: bool = False,
        font_color: str = None,
        highlight_color: str = None
) -> str:
    """
    设置Word文档中指定段落的文本格式
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param paragraph_index: 段落索引 (从0开始计数)
    :param font_name: 字体名称 (可选)
    :param font_size: 字体大小 (点数，可选)
    :param bold: 是否加粗
    :param italic: 是否斜体
    :param underline: 是否下划线
    :param font_color: 字体颜色 (十六进制RGB格式，如"#FF0000"表示红色，可选)
    :param highlight_color: 突出显示颜色 (可选，有效值: "yellow", "green", "blue", "red", "pink", "turquoise", "violet", "darkblue", "teal", "darkred", "darkgreen")
    :
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"错误: 文件 {file_path} 不存在"})

    # 定义高亮颜色映射 (用于XML着色)
    highlight_color_map = {
        "yellow": "FFFF00",
        "green": "00FF00",
        "blue": "0000FF",
        "red": "FF0000",
        "pink": "FFC0CB",
        "turquoise": "40E0D0",
        "violet": "EE82EE",
        "darkblue": "00008B",
        "teal": "008080",
        "darkred": "8B0000",
        "darkgreen": "006400"
    }

    # 校验高亮颜色
    if highlight_color and highlight_color.lower() not in highlight_color_map:
        return response_handler({"status": "error",
                                 "message": f"错误: 不支持的高亮颜色 '{highlight_color}'，可选值为: {', '.join(highlight_color_map.keys())}"})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"})

        # 获取指定的段落
        paragraph = doc.paragraphs[paragraph_index]

        # 检查段落是否有内容
        if not paragraph.text.strip():
            return response_handler(
                {"status": "warning", "message": f"警告: 段落 {paragraph_index + 1} 为空或只包含空白字符，无法设置格式"})

        # 检查段落是否有run，如果没有，添加一个run
        if len(paragraph.runs) == 0:
            # 保存原始文本
            original_text = paragraph.text
            # 清空段落
            for child in list(paragraph._element):
                paragraph._element.remove(child)
            # 添加新run
            paragraph.add_run(original_text)

        # 应用格式设置
        for run in paragraph.runs:
            if font_name:
                # 设置西文字体名
                run.font.name = font_name
                # 设置中文字体名
                run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), font_name)

            if font_size:
                run.font.size = Pt(font_size)

            run.font.bold = bold
            run.font.italic = italic
            run.font.underline = underline

            # 设置字体颜色
            if font_color:
                try:
                    # 移除#号并统一处理缩写格式
                    font_color = font_color.lstrip('#').upper()

                    # 处理缩写格式（如 #FFF -> FFFFFF）
                    if len(font_color) == 3:
                        font_color = ''.join([c * 2 for c in font_color])

                    # 验证是否为有效的6位十六进制
                    if len(font_color) != 6:
                        raise ValueError("颜色格式错误")

                    r = int(font_color[0:2], 16)
                    g = int(font_color[2:4], 16)
                    b = int(font_color[4:6], 16)
                    run.font.color.rgb = RGBColor(r, g, b)
                except ValueError:
                    return response_handler({"status": "error",
                                             "message": f"错误: 无效的字体颜色格式 '{font_color}'，请使用十六进制RGB格式，如 '#FF0000'"})

            # 设置高亮颜色（通过XML方式）
            if highlight_color:
                shading_elm = OxmlElement('w:shd')
                color_value = highlight_color_map[highlight_color.lower()]
                shading_elm.set(qn('w:fill'), color_value)
                run._element.get_or_add_rPr().append(shading_elm)

        # 保存文档
        doc.save(file_path)

        return response_handler(
            {"status": "success", "message": f"成功设置文档 {os.path.basename(file_path)} 第 {paragraph_index + 1} 段落的格式"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"设置Word文档格式时出错: {str(e)}"})


@mcp.tool()
def set_paragraph_spacing(
        file_path: str,
        paragraph_index: int,
        before_spacing: float = None,
        after_spacing: float = None,
        line_spacing: float = None,
        line_spacing_rule: str = "multiple"
) -> str:
    """
    设置Word文档中指定段落的间距
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param paragraph_index: 段落索引 (从0开始计数)
    :param before_spacing: 段前间距 (磅，可选)
    :param after_spacing: 段后间距 (磅，可选)
    :param line_spacing: 行间距值 (当line_spacing_rule为"multiple"时为倍数，为"exact"时为磅值)
    :param line_spacing_rule: 行间距规则，可选值: "multiple"(倍数), "exact"(固定值), "atLeast"(最小值)
    :
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"错误: 文件 {file_path} 不存在"})

    # 映射行间距规则
    spacing_rule_map = {
        "multiple": WD_LINE_SPACING.MULTIPLE,
        "exact": WD_LINE_SPACING.EXACTLY,
        "atLeast": WD_LINE_SPACING.AT_LEAST
    }

    if line_spacing_rule not in spacing_rule_map:
        return response_handler(
            {"status": "error", "message": f"错误: 无效的行间距规则 '{line_spacing_rule}'，可选值为: multiple, exact, atLeast"})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"错误: 无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落"})

        # 获取指定的段落
        paragraph = doc.paragraphs[paragraph_index]

        # 设置段前间距
        if before_spacing is not None:
            paragraph.paragraph_format.space_before = Pt(before_spacing)

        # 设置段后间距
        if after_spacing is not None:
            paragraph.paragraph_format.space_after = Pt(after_spacing)

        # 设置行间距
        if line_spacing is not None:
            # 设置行间距规则
            paragraph.paragraph_format.line_spacing_rule = spacing_rule_map[line_spacing_rule]

            # 根据规则设置行间距值
            if line_spacing_rule == "multiple":
                paragraph.paragraph_format.line_spacing = line_spacing
            else:
                paragraph.paragraph_format.line_spacing = Pt(line_spacing)

        # 保存文档
        doc.save(file_path)

        return response_handler(
            {"status": "success", "message": f"成功设置文档 {os.path.basename(file_path)} 第 {paragraph_index + 1} 段落的间距"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"设置段落间距时出错: {str(e)}"})


@mcp.tool()
def insert_image(
        file_path: str,
        image_path: str,
        width: float = None,
        height: float = None,
        after_paragraph: int = -1
) -> str:
    """
    在Word文档中插入图片
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param image_path: 图片文件的完整路径或相对于输出目录的路径，支持本地图片
    :param width: 图片宽度（厘米，可选，如果不指定则保持原始比例）
    :param height: 图片高度（厘米，可选，如果不指定则保持原始比例）
    :param after_paragraph: 在指定段落后插入图片，-1表示文档末尾
    :
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"错误: 文件 {file_path} 不存在"})

    # 处理图片路径，同样支持相对路径
    if not os.path.isabs(image_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        image_path = os.path.join(base_path, image_path)

    # 检查图片文件是否存在
    if not os.path.exists(image_path):
        return response_handler({"status": "error", "message": f"错误: 图片文件 {image_path} 不存在"})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查指定段落是否有效
        if after_paragraph >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"错误: 无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落"})

        # 在指定位置插入图片
        if after_paragraph == -1:
            # 在文档末尾插入图片
            paragraph = doc.add_paragraph()
        else:
            # 在指定段落后插入新段落，然后插入图片
            paragraph = doc.paragraphs[after_paragraph]

        # 设置图片尺寸
        if width and height:
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width), height=Cm(height))
        elif width:
            run = paragraph.add_run()
            run.add_picture(image_path, width=Cm(width))
        elif height:
            run = paragraph.add_run()
            run.add_picture(image_path, height=Cm(height))
        else:
            run = paragraph.add_run()
            run.add_picture(image_path)

        # 保存文档
        doc.save(file_path)

        return response_handler({"status": "success",
                                 "message": f"成功在文档 {os.path.basename(file_path)} 中插入图片 {os.path.basename(image_path)}"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"插入图片时出错: {str(e)}"})


@mcp.tool()
def insert_table(
        file_path: str,
        rows: int,
        cols: int,
        data: List[List[str]] = None,
        after_paragraph: int = -1,
        style: str = "Table Grid"
) -> str:
    """
    在Word文档中插入表格
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param rows: 表格行数
    :param cols: 表格列数
    :param data: 表格内容，二维数组，每个元素对应一个单元格的内容（可选）
    :param after_paragraph: 在指定段落后插入表格，-1表示文档末尾
    :param style: 表格样式，默认为"Table Grid"
    """
    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    if rows <= 0 or cols <= 0:
        return response_handler({"status": "error", "message": "表格行数和列数必须大于0", "data": None})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查指定段落是否有效
        if after_paragraph >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落",
                 "data": None})

        # 在指定位置插入表格
        if after_paragraph == -1:
            # 在文档末尾插入表格
            table = doc.add_table(rows=rows, cols=cols)
        else:
            # 获取指定段落的位置
            paragraph = doc.paragraphs[after_paragraph]
            # 在段落后插入表格
            table = doc.add_table(rows=rows, cols=cols)
            # 移动表格到指定段落后
            paragraph._p.addnext(table._tbl)

        # 设置表格样式
        table.style = style

        # 如果提供了数据，填充表格内容
        if data:
            for i, row_data in enumerate(data):
                if i < rows:  # 确保不超出表格行数
                    for j, cell_data in enumerate(row_data):
                        if j < cols:  # 确保不超出表格列数
                            table.cell(i, j).text = str(cell_data)

        # 保存文档
        doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功在文档 {os.path.basename(file_path)} 中插入 {rows}x{cols} 的表格",
            "data": {
                "file_path": file_path,
                "rows": rows,
                "cols": cols,
                "table_index": len(doc.tables) - 1
            }
        })
    except Exception as e:
        return response_handler({"status": "error", "message": f"插入表格时出错: {str(e)}", "data": None})


@mcp.tool()
def edit_table_cell(
        file_path: str,
        table_index: int,
        row: int,
        col: int,
        text: str
) -> str:
    """
    编辑Word文档中表格的单元格内容
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param table_index: 表格索引 (从0开始计数)
    :param row: 行索引 (从0开始计数)
    :param col: 列索引 (从0开始计数)
    :param text: 单元格内容
    """

    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查表格索引是否有效
        if table_index < 0 or table_index >= len(doc.tables):
            return response_handler(
                {"status": "error", "message": f"无效的表格索引 {table_index}，文档共有 {len(doc.tables)} 个表格", "data": None})

        # 获取指定的表格
        table = doc.tables[table_index]

        # 检查行索引是否有效
        if row < 0 or row >= len(table.rows):
            return response_handler(
                {"status": "error", "message": f"无效的行索引 {row}，表格共有 {len(table.rows)} 行", "data": None})

        # 检查列索引是否有效
        if col < 0 or col >= len(table.columns):
            return response_handler(
                {"status": "error", "message": f"无效的列索引 {col}，表格共有 {len(table.columns)} 列", "data": None})

        # 编辑单元格内容
        table.cell(row, col).text = text

        # 保存文档
        doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功编辑文档 {os.path.basename(file_path)} 中第 {table_index + 1} 个表格的单元格 ({row + 1},{col + 1})",
            "data": {
                "file_path": file_path,
                "table_index": table_index,
                "row": row,
                "col": col,
                "new_text": text
            }
        })
    except Exception as e:
        return response_handler({"status": "error", "message": f"编辑表格单元格时出错: {str(e)}", "data": None})


@mcp.tool()
def save_document_as(file_path: str, output_format: str = "docx", new_filename: str = None) -> str:
    """
    将Word文档保存为指定格式
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param output_format: 输出格式，可选值: "docx", "doc", "pdf", "txt", "html"
    :param new_filename: 新文件名(不含扩展名)，如果不提供则使用原文件名
    """

    # 检查格式是否支持
    supported_formats = ["docx", "pdf", "txt", "html"]
    if output_format.lower() not in supported_formats:
        return response_handler(
            {"status": "error", "message": f"不支持的输出格式 '{output_format}'，可选值为: {', '.join(supported_formats)}",
             "data": None})

    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    try:
        # 构建新文件路径
        original_basename = os.path.splitext(os.path.basename(file_path))[0]
        output_dirname = os.path.dirname(file_path)

        # 如果提供了新文件名，则使用新文件名
        if new_filename:
            output_basename = new_filename
        else:
            output_basename = original_basename

        # 创建新文件的完整路径
        output_path = os.path.join(output_dirname, f"{output_basename}.{output_format}")

        # 根据输出格式选择不同的处理方式
        if output_format.lower() == "pdf":
            # 使用 reportlab 生成 PDF
            doc = Document(file_path)
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            # 创建 PDF
            pdf = SimpleDocTemplate(output_path)
            styles = getSampleStyleSheet()
            styles['Normal'].fontName = 'SimHei'  # 设置中文字体, 防止乱码
            elements = []

            for para in doc.paragraphs:
                if para.text.strip():
                    elements.append(Paragraph(para.text, styles['Normal']))
                    elements.append(Spacer(1, 0.2 * inch))

            pdf.build(elements)

            return response_handler({
                "status": "success",
                "message": f"成功将文档导出为PDF: {os.path.basename(output_path)}",
                "data": {
                    "original_file": file_path,
                    "pdf_file": output_path
                }
            })

        elif output_format.lower() == "docx":
            # 使用 python-docx 保存为 DOCX
            doc = Document(file_path)
            doc.save(output_path)

            return response_handler({
                "status": "success",
                "message": f"成功将文档保存为 DOCX 格式: {os.path.basename(output_path)}",
                "data": {
                    "original_file": file_path,
                    "new_file": output_path,
                    "format": "docx"
                }
            })

        elif output_format.lower() == "txt":
            # 将文档转换为纯文本
            doc = Document(file_path)
            text_content = "\n\n".join([para.text for para in doc.paragraphs if para.text.strip()])

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text_content)

            return response_handler({
                "status": "success",
                "message": f"成功将文档保存为文本格式: {os.path.basename(output_path)}",
                "data": {
                    "original_file": file_path,
                    "new_file": output_path,
                    "format": "txt"
                }
            })

        elif output_format.lower() == "html":
            # 将文档转换为 HTML
            doc = Document(file_path)
            html_content = "<html><body>\n"
            for para in doc.paragraphs:
                if para.text.strip():
                    html_content += f"<p>{para.text}</p>\n"
            html_content += "</body></html>"

            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            return response_handler({
                "status": "success",
                "message": f"成功将文档保存为HTML格式: {os.path.basename(output_path)}",
                "data": {
                    "original_file": file_path,
                    "new_file": output_path,
                    "format": "html"
                }
            })

    except Exception as e:
        return response_handler({"status": "error", "message": f"保存文档时出错: {str(e)}", "data": None})


@mcp.tool()
def convert_to_docx(file_path: str, new_filename: str = None) -> str:
    """
    将支持的文件格式转换为DOCX格式
    :param file_path: 原始文件的完整路径或相对于工作目录的路径
    :param new_filename: 新文件名(不含扩展名)，如果不提供则使用原文件名
    """
    # 支持的输入格式
    supported_formats = ["pdf", "txt", "html", "docx"]

    try:
        # 处理文件路径
        if not os.path.isabs(file_path):
            base_path = os.environ.get('WORD_MCP_PATH', os.path.join(os.path.expanduser('~'), '桌面'))
            file_path = os.path.join(base_path, file_path)

        if not os.path.exists(file_path):
            return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

        # 获取文件信息
        file_ext = os.path.splitext(file_path)[1][1:].lower()
        original_basename = os.path.splitext(os.path.basename(file_path))[0]

        # 检查格式支持
        if file_ext not in supported_formats:
            return response_handler({
                "status": "error",
                "message": f"不支持的文件格式 '{file_ext}'，支持格式: {', '.join(supported_formats)}",
                "data": None
            })

        # 如果是docx直接返回原文件
        if file_ext == "docx":
            return response_handler({
                "status": "success",
                "message": "文件已经是DOCX格式",
                "data": {
                    "original_file": file_path,
                    "new_file": file_path,
                    "format": "docx"
                }
            })

        # 构建输出路径
        output_dir = os.path.dirname(file_path)
        output_basename = new_filename if new_filename else original_basename
        output_path = os.path.join(output_dir, f"{output_basename}.docx")

        if file_ext == "pdf":
            try:
                # 使用 PyPDF2 提取 PDF 文本
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                print(f"提取的文本长度: {len(text)}")  # 调试信息

                # 将文本写入到新的 docx 文档中
                doc = Document()
                paragraphs = text.split("\n")
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
                print(f"即将保存到: {output_path}")  # 调试信息
                doc.save(output_path)
            except Exception as e:
                return response_handler({
                    "status": "error",
                    "message": f"PDF转DOCX失败: {str(e)}",
                    "data": None
                })

        elif file_ext == "txt":
            try:
                # 使用 python-docx 创建新文档
                doc = Document()
                with open(file_path, 'r', encoding='utf-8') as f:
                    for line in f:
                        # 按换行符分割段落
                        doc.add_paragraph(line.strip())
                doc.save(output_path)
            except Exception as e:
                return response_handler({
                    "status": "error",
                    "message": f"TXT转DOCX失败: {str(e)}",
                    "data": None
                })

        elif file_ext == "html":
            try:
                # 使用正则表达式提取 HTML 中的文本
                with open(file_path, 'r', encoding='utf-8') as f:
                    html_content = f.read()
                    text = re.sub(r'<[^>]+>', '', html_content)
                # 将文本写入到新的 docx 文档中
                doc = Document()
                doc.add_paragraph(text)
                doc.save(output_path)
            except Exception as e:
                return response_handler({
                    "status": "error",
                    "message": f"HTML转DOCX失败: {str(e)}",
                    "data": None
                })

        return response_handler({
            "status": "success",
            "message": f"成功转换文件为DOCX格式: {os.path.basename(output_path)}",
            "data": {
                "original_file": file_path,
                "new_file": output_path,
                "format": "docx"
            }
        })

    except Exception as e:
        return response_handler({
            "status": "error",
            "message": f"文件转换失败: {str(e)}",
            "data": None
        })


@mcp.tool()
def close_document(file_path: str, save_changes: bool = True) -> str:
    """
    关闭Word文档，可选是否保存更改
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param save_changes: 是否保存更改，默认为True
    """

    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    try:
        # 使用 python-docx 读取文档
        doc = Document(file_path)

        # 如果需要保存更改，则保存文档
        if save_changes:
            doc.save(file_path)

        # python-docx 没有明确的关闭方法，垃圾收集器会处理
        return response_handler({
            "status": "success",
            "message": f"成功关闭文档: {os.path.basename(file_path)}" + (" 并保存更改" if save_changes else ""),
            "data": {
                "file_path": file_path,
                "saved": save_changes
            }
        })

    except Exception as e:
        return response_handler({"status": "error", "message": f"关闭文档时出错: {str(e)}", "data": None})


@mcp.tool()
def edit_paragraph_in_document(
        file_path: str,
        paragraph_index: int,
        new_text: str,
        save: bool = True
) -> str:
    """
    编辑Word文档中指定段落的文本内容
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param paragraph_index: 段落索引 (从0开始计数)
    :param new_text: 新的文本内容
    :param save: 是否保存更改，默认为True
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落",
                 "data": None})

        # 获取并编辑指定的段落
        paragraph = doc.paragraphs[paragraph_index]

        # 保存原始样式和格式
        original_style = paragraph.style
        original_alignment = paragraph.alignment

        # 更简单安全的替换方法：清除所有runs并添加新文本
        for run in paragraph.runs:
            run.clear()

        # 清空所有runs后，确保段落内容被清除
        if paragraph.runs:
            # 如果仍有runs，直接重新设置text属性
            paragraph.text = ""

        # 添加新内容
        run = paragraph.add_run(new_text)

        # 恢复原始样式和格式
        paragraph.style = original_style
        paragraph.alignment = original_alignment

        # 保存文档
        if save:
            doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功编辑文档 {os.path.basename(file_path)} 第 {paragraph_index + 1} 段落的内容",
            "data": {
                "file_path": file_path,
                "paragraph_index": paragraph_index,
                "new_text": new_text,
                "saved": save
            }
        })
    except Exception as e:
        return response_handler({"status": "error", "message": f"编辑Word文档内容时出错: {str(e)}", "data": None})


@mcp.tool()
def find_and_replace_text(
        file_path: str,
        find_text: str,
        replace_text: str,
        match_case: bool = False,
        match_whole_word: bool = False,
        save: bool = True
) -> str:
    """
    在Word文档中查找并替换文本
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param find_text: 要查找的文本
    :param replace_text: 替换为的文本
    :param match_case: 是否区分大小写，默认为False
    :param match_whole_word: 是否匹配整个单词，默认为False
    :param save: 是否保存更改，默认为True
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    try:
        # 使用python-docx的方式（更可靠）
        doc = Document(file_path)
        replace_count = 0

        # 遍历所有段落和所有run
        for paragraph in doc.paragraphs:
            # 获取段落的完整文本
            full_text = paragraph.text
            if not match_case:
                search_text = find_text.lower()
                full_text_lower = full_text.lower()
            else:
                search_text = find_text
                full_text_lower = full_text

            # 如果段落中包含要查找的文本
            if search_text in full_text_lower:
                # 清除所有runs
                for run in paragraph.runs:
                    run.clear()

                # 如果需要区分大小写
                if match_case:
                    # 直接替换
                    new_text = full_text.replace(find_text, replace_text)
                    replace_count += full_text.count(find_text)
                else:
                    # 不区分大小写的替换
                    current_pos = 0
                    new_text = ""
                    while True:
                        # 在剩余文本中查找目标字符串
                        pos = full_text_lower[current_pos:].find(search_text)
                        if pos == -1:
                            # 没有找到更多匹配，添加剩余文本
                            new_text += full_text[current_pos:]
                            break

                        # 添加匹配位置之前的文本
                        new_text += full_text[current_pos:current_pos + pos]
                        # 添加替换文本
                        new_text += replace_text
                        # 更新位置
                        current_pos += pos + len(find_text)
                        replace_count += 1

                # 添加新的run，包含替换后的文本
                paragraph.add_run(new_text)

        # 遍历所有表格单元格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        # 获取段落的完整文本
                        full_text = paragraph.text
                        if not match_case:
                            search_text = find_text.lower()
                            full_text_lower = full_text.lower()
                        else:
                            search_text = find_text
                            full_text_lower = full_text

                        # 如果段落中包含要查找的文本
                        if search_text in full_text_lower:
                            # 清除所有runs
                            for run in paragraph.runs:
                                run.clear()

                            # 如果需要区分大小写
                            if match_case:
                                # 直接替换
                                new_text = full_text.replace(find_text, replace_text)
                                replace_count += full_text.count(find_text)
                            else:
                                # 不区分大小写的替换
                                current_pos = 0
                                new_text = ""
                                while True:
                                    # 在剩余文本中查找目标字符串
                                    pos = full_text_lower[current_pos:].find(search_text)
                                    if pos == -1:
                                        # 没有找到更多匹配，添加剩余文本
                                        new_text += full_text[current_pos:]
                                        break

                                    # 添加匹配位置之前的文本
                                    new_text += full_text[current_pos:current_pos + pos]
                                    # 添加替换文本
                                    new_text += replace_text
                                    # 更新位置
                                    current_pos += pos + len(find_text)
                                    replace_count += 1

                            # 添加新的run，包含替换后的文本
                            paragraph.add_run(new_text)

        # 保存文档
        if save:
            doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功在文档 {os.path.basename(file_path)} 中替换了 {replace_count} 处文本",
            "data": {
                "file_path": file_path,
                "find_text": find_text,
                "replace_text": replace_text,
                "replace_count": replace_count,
                "saved": save
            }
        })

    except Exception as e:
        return response_handler({"status": "error", "message": f"在Word文档中查找替换文本时出错: {str(e)}", "data": None})


@mcp.tool()
def delete_paragraph(
        file_path: str,
        paragraph_index: int,
        save: bool = True
) -> str:
    """
    删除Word文档中指定的段落
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param paragraph_index: 要删除的段落索引 (从0开始计数)
    :param save: 是否保存更改，默认为True
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查段落索引是否有效
        if paragraph_index < 0 or paragraph_index >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"无效的段落索引 {paragraph_index}，文档共有 {len(doc.paragraphs)} 个段落",
                 "data": None})

        # 获取要删除的段落
        paragraph = doc.paragraphs[paragraph_index]

        # 删除段落
        p = paragraph._element
        p.getparent().remove(p)

        # 删除对象的引用
        paragraph._p = None
        paragraph._element = None

        # 保存文档
        if save:
            doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功从文档 {os.path.basename(file_path)} 中删除第 {paragraph_index + 1} 段落",
            "data": {
                "file_path": file_path,
                "paragraph_index": paragraph_index,
                "saved": save
            }
        })
    except Exception as e:
        return response_handler({"status": "error", "message": f"删除Word文档段落时出错: {str(e)}", "data": None})


@mcp.tool()
def insert_table_of_contents(
        file_path: str,
        title: str = "目录",
        levels: int = 3,
        after_paragraph: int = 0
) -> str:
    """
    在Word文档中插入目录
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param title: 目录标题
    :param levels: 目录级别数 (1-9)
    :param after_paragraph: 在指定段落后插入目录，默认为文档开头第一段后
    """

    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    if levels < 1 or levels > 9:
        return response_handler({"status": "error", "message": "目录级别数必须在1至9之间", "data": None})

    try:
        doc = Document(file_path)

        # 检查指定段落是否有效
        if after_paragraph >= len(doc.paragraphs):
            return response_handler(
                {"status": "error", "message": f"无效的段落索引 {after_paragraph}，文档共有 {len(doc.paragraphs)} 个段落",
                 "data": None})

        # 在指定位置插入目录标题
        if after_paragraph == 0:
            # 在文档开头插入
            if title:
                heading_para = doc.add_paragraph(title, style="Heading 1")
                if len(doc.paragraphs) > 1:
                    first_para = doc.paragraphs[1]._p
                    heading_para._p.addnext(first_para)
        else:
            # 在指定段落后插入
            paragraph = doc.paragraphs[after_paragraph]
            if title:
                new_para = doc.add_paragraph()
                paragraph._p.addnext(new_para._p)
                new_para.text = title
                new_para.style = "Heading 1"

        # 创建目录字段
        toc_para = doc.add_paragraph()
        if title:
            if after_paragraph == 0:
                if len(doc.paragraphs) > 1:
                    doc.paragraphs[1]._p.addnext(toc_para._p)
                else:
                    doc.add_paragraph(toc_para.text)
            else:
                doc.paragraphs[after_paragraph + 1]._p.addnext(toc_para._p)
        else:
            paragraph._p.addnext(toc_para._p)

        toc_run = toc_para.add_run()

        # 添加目录字段XML
        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'begin')
        toc_run._r.append(fldChar)

        instrText = OxmlElement('w:instrText')
        instrText.text = f' TOC \\o "1-{levels}" \\h '
        toc_run._r.append(instrText)

        fldChar = OxmlElement('w:fldChar')
        fldChar.set(qn('w:fldCharType'), 'end')
        toc_run._r.append(fldChar)

        # 保存文档
        doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功在文档 {os.path.basename(file_path)} 中插入目录（需要在Word中手动更新）",
            "data": {
                "file_path": file_path,
                "title": title,
                "levels": levels,
                "after_paragraph": after_paragraph,
                "requires_manual_update": True
            }
        })

    except Exception as e:
        return response_handler({"status": "error", "message": f"插入目录时出错: {str(e)}", "data": None})


@mcp.tool()
def add_header_footer(
        file_path: str,
        header_text: str = None,
        footer_text: str = None,
        page_numbers: bool = False
) -> str:
    """
    为Word文档添加页眉和页脚
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param header_text: 页眉文本（可选）
    :param footer_text: 页脚文本（可选）
    :param page_numbers: 是否在页脚添加页码
    """

    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')
        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    # 检查是否提供了有效的参数
    if header_text is None and footer_text is None and not page_numbers:
        return response_handler({"status": "error", "message": "请至少提供页眉文本、页脚文本或启用页码", "data": None})

    try:
        doc = Document(file_path)

        # 获取所有节
        sections = doc.sections

        # 为每个节添加页眉页脚
        for section in sections:
            # 添加页眉
            if header_text:
                header = section.header
                header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                header_para.text = header_text
                header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in header_para.runs:
                    run.font.size = Pt(12)

            # 添加页脚
            if footer_text or page_numbers:
                footer = section.footer
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()

                if footer_text:
                    footer_para.text = footer_text
                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in footer_para.runs:
                        run.font.size = Pt(12)

                # 添加页码（python-docx对页码支持有限）
                if page_numbers:
                    run = footer_para.add_run()

                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'begin')
                    run._r.append(fldChar)

                    instrText = OxmlElement('w:instrText')
                    instrText.text = ' PAGE '
                    run._r.append(instrText)

                    fldChar = OxmlElement('w:fldChar')
                    fldChar.set(qn('w:fldCharType'), 'end')
                    run._r.append(fldChar)

                    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 保存文档
        doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功为文档 {os.path.basename(file_path)} 添加页眉页脚",
            "data": {
                "file_path": file_path,
                "header_text": header_text,
                "footer_text": footer_text,
                "page_numbers": page_numbers
            }
        })

    except Exception as e:
        return response_handler({"status": "error", "message": f"添加页眉页脚时出错: {str(e)}", "data": None})


@mcp.tool()
def set_page_layout(
        file_path: str,
        orientation: str = None,
        page_width: float = None,
        page_height: float = None,
        left_margin: float = None,
        right_margin: float = None,
        top_margin: float = None,
        bottom_margin: float = None,
        section_index: int = 0
) -> str:
    """
    设置Word文档的页面布局
    :param file_path: Word文档的完整路径或相对于输出目录的路径
    :param orientation: 页面方向，可选值: "portrait"(纵向), "landscape"(横向)
    :param page_width: 页面宽度（厘米，自定义纸张尺寸时使用）
    :param page_height: 页面高度（厘米，自定义纸张尺寸时使用）
    :param left_margin: 左边距（厘米）
    :param right_margin: 右边距（厘米）
    :param top_margin: 上边距（厘米）
    :param bottom_margin: 下边距（厘米）
    :param section_index: 节索引，默认为0（第一节）
    """

    if not os.path.isabs(file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

    # 校验方向参数
    orientation_map = {
        "portrait": WD_ORIENTATION.PORTRAIT,
        "landscape": WD_ORIENTATION.LANDSCAPE
    }

    if orientation and orientation.lower() not in orientation_map:
        return response_handler(
            {"status": "error", "message": f"无效的页面方向 '{orientation}'，可选值为: portrait, landscape", "data": None})

    try:
        # 打开Word文档
        doc = Document(file_path)

        # 检查节索引是否有效
        if section_index < 0 or section_index >= len(doc.sections):
            return response_handler(
                {"status": "error", "message": f"无效的节索引 {section_index}，文档共有 {len(doc.sections)} 节", "data": None})

        # 获取指定的节
        section = doc.sections[section_index]

        # 设置页面方向
        if orientation:
            section.orientation = orientation_map[orientation.lower()]

        # 设置页面尺寸
        if page_width and page_height:
            section.page_width = Cm(page_width)
            section.page_height = Cm(page_height)

        # 设置页边距
        if left_margin is not None:
            section.left_margin = Cm(left_margin)

        if right_margin is not None:
            section.right_margin = Cm(right_margin)

        if top_margin is not None:
            section.top_margin = Cm(top_margin)

        if bottom_margin is not None:
            section.bottom_margin = Cm(bottom_margin)

        # 保存文档
        doc.save(file_path)

        return response_handler({
            "status": "success",
            "message": f"成功设置文档 {os.path.basename(file_path)} 第 {section_index + 1} 节的页面布局",
            "data": {
                "file_path": file_path,
                "section_index": section_index,
                "orientation": orientation,
                "page_width": page_width,
                "page_height": page_height,
                "left_margin": left_margin,
                "right_margin": right_margin,
                "top_margin": top_margin,
                "bottom_margin": bottom_margin
            }
        })
    except Exception as e:
        return response_handler({"status": "error", "message": f"设置页面布局时出错: {str(e)}", "data": None})


@mcp.tool()
def merge_documents(
        main_file_path: str,
        files_to_merge: List[str]
) -> str:
    """
    合并多个Word文档
    :param main_file_path: 主文档的完整路径或相对于输出目录的路径（合并后的文档将保存为该文件）
    :param files_to_merge: 要合并的文档路径列表
    """

    if not os.path.isabs(main_file_path):

        base_path = os.environ.get('WORD_MCP_PATH')
        if not base_path:
            base_path = os.path.join(os.path.expanduser('~'), '桌面')

        main_file_path = os.path.join(base_path, main_file_path)

    if not files_to_merge:
        return response_handler({"status": "error", "message": "请提供至少一个要合并的文档", "data": None})

    # 处理文件路径
    processed_files = []
    for file_path in files_to_merge:
        if not os.path.isabs(file_path):

            base_path = os.environ.get('WORD_MCP_PATH')
            if not base_path:
                base_path = os.path.join(os.path.expanduser('~'), '桌面')

            file_path = os.path.join(base_path, file_path)

        if not os.path.exists(file_path):
            return response_handler({"status": "error", "message": f"文件 {file_path} 不存在", "data": None})

        processed_files.append(file_path)

    try:
        # 尝试使用Word COM对象合并文档（功能最完整）
        try:
            import win32com.client
            import pythoncom

            pythoncom.CoInitialize()
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # 检查主文档是否存在，如果不存在则创建
            if not os.path.exists(main_file_path):
                doc = word.Documents.Add()
                doc.SaveAs(main_file_path)
            else:
                doc = word.Documents.Open(main_file_path)

            # 记录成功合并的文档数量
            merged_count = 0

            # 合并每个文档
            for file_path in processed_files:
                # 将光标移动到文档末尾
                word.Selection.EndKey(Unit=6)  # 6表示wdStory，即整个文档

                # 插入分节符
                if merged_count > 0:
                    word.Selection.InsertBreak(Type=2)  # 2表示wdSectionBreakNextPage

                # 插入文档内容
                word.Selection.InsertFile(file_path)
                merged_count += 1

            # 保存并关闭
            doc.Save()
            doc.Close()
            word.Quit()

            pythoncom.CoUninitialize()

            return response_handler({
                "status": "success",
                "message": f"成功将 {merged_count} 个文档合并到 {os.path.basename(main_file_path)}",
                "data": {
                    "main_file": main_file_path,
                    "merged_files": processed_files,
                    "merged_count": merged_count
                }
            })

        except ImportError:
            # 使用python-docx方式合并文档（功能受限）
            if os.path.exists(main_file_path):
                main_doc = Document(main_file_path)
            else:
                main_doc = Document()

            # 记录成功合并的文档数量
            merged_count = 0

            # 合并每个文档
            for file_path in processed_files:
                # 打开要合并的文档
                doc_to_merge = Document(file_path)

                # 插入分节符（如果不是第一个文档）
                if merged_count > 0:
                    main_doc.add_section()

                # 复制所有段落
                for paragraph in doc_to_merge.paragraphs:
                    new_paragraph = main_doc.add_paragraph()
                    for run in paragraph.runs:
                        new_run = new_paragraph.add_run(run.text)
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        if run.font.size:
                            new_run.font.size = run.font.size
                        if run.font.name:
                            new_run.font.name = run.font.name

                    new_paragraph.style = paragraph.style
                    new_paragraph.alignment = paragraph.alignment

                # 复制所有表格
                for table in doc_to_merge.tables:
                    new_table = main_doc.add_table(rows=len(table.rows), cols=len(table.columns))
                    new_table.style = table.style

                    for i, row in enumerate(table.rows):
                        for j, cell in enumerate(row.cells):
                            if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                                new_table.rows[i].cells[j].text = cell.text

                merged_count += 1

            # 保存合并后的文档
            main_doc.save(main_file_path)

            return response_handler({
                "status": "success",
                "message": f"成功将 {merged_count} 个文档合并到 {os.path.basename(main_file_path)}",
                "data": {
                    "main_file": main_file_path,
                    "merged_files": processed_files,
                    "merged_count": merged_count
                }
            })

    except Exception as e:
        return response_handler({"status": "error", "message": f"合并文档时出错: {str(e)}", "data": None})


@mcp.tool()
def complex_query(file_path: str, query: str):
    """
    复杂查询Word文档
    :param file_path: 文档路径
    :param query: 查询表达式（支持格式）：
        - regex:pattern      正则表达式匹配
        - keyword:text       关键字匹配（完整词）
        - contains:text      包含匹配
        - tables             统计表格数量
        - images             统计图片数量
        - paragraphs         统计段落数量
    :return: 返回包含统计结果和匹配详情的JSON
    """
    import os
    import re
    from docx import Document

    # 路径标准化处理
    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH', os.path.expanduser('~/桌面'))
        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": "文件不存在"})

    try:
        doc = Document(file_path)
    except Exception as e:
        return response_handler({"status": "error", "message": f"文档解析失败: {str(e)}"})

    # 解析查询指令
    query_type = "raw"
    search_pattern = query
    flags = re.IGNORECASE  # 默认不区分大小写

    if ':' in query:
        query_type, search_pattern = query.split(':', 1)
        query_type = query_type.lower().strip()

    result = {
        "total": 0,
        "details": [],
        "elements": {}
    }

    try:
        # 处理不同查询类型
        if query_type == 'regex':
            pattern = re.compile(search_pattern, flags)
        elif query_type in ('keyword', 'contains'):
            escaped = re.escape(search_pattern)
            pattern = re.compile(rf'\b{escaped}\b' if query_type == 'keyword' else escaped, flags)
        else:  # 统计类查询
            pass

        # 文本内容查询逻辑
        if query_type in ('regex', 'keyword', 'contains'):
            # 遍历段落
            for para_idx, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if matches := pattern.findall(text):
                    result["total"] += len(matches)
                    result["details"].append({
                        "type": "paragraph",
                        "index": para_idx,
                        "text_snippet": text[:50] + "..." if len(text) > 50 else text,
                        "matches": matches
                    })

            # 遍历表格
            for table_idx, table in enumerate(doc.tables):
                for row_idx, row in enumerate(table.rows):
                    for cell_idx, cell in enumerate(row.cells):
                        text = cell.text.strip()
                        if matches := pattern.findall(text):
                            result["total"] += len(matches)
                            result["details"].append({
                                "type": "table",
                                "position": f"Table-{table_idx} Cell({row_idx},{cell_idx})",
                                "text_snippet": text[:50] + "..." if len(text) > 50 else text,
                                "matches": matches
                            })

        # 元素统计查询逻辑
        elif query_type == 'raw':
            if query == 'tables':
                result["elements"]["tables"] = len(doc.tables)
            elif query == 'images':
                result["elements"]["images"] = len(doc.inline_shapes)
            elif query == 'paragraphs':
                result["elements"]["paragraphs"] = len(doc.paragraphs)
            else:
                return response_handler({"status": "error", "message": "不支持的查询类型"})

        return response_handler({
            "status": "success",
            "data": {
                "query_type": query_type,
                **result
            }
        })

    except re.error as e:
        return response_handler({"status": "error", "message": f"正则表达式错误: {str(e)}"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"查询执行失败: {str(e)}"})


@mcp.tool()
def complex_replace(file_path: str, replace: str, output_path: str = None):
    """
    替换 Word 文档内容（保留原始格式）
    :param file_path: 原文件路径
    :param replace: 替换指令（支持格式）：
        - regex:pattern=replacement  正则表达式替换
        - keyword:old=new            完整词替换
        - contains:old=new          包含文本替换
    :param output_path: 新文件保存路径（默认添加 _modified 后缀）
    :return: 包含替换统计和新文件路径的 JSON
    """

    # 路径处理
    if not os.path.isabs(file_path):
        base_path = os.environ.get('WORD_MCP_PATH', os.path.expanduser('~/桌面'))
        file_path = os.path.join(base_path, file_path)

    if not os.path.exists(file_path):
        return response_handler({"status": "error", "message": "文件不存在"})

    # 处理输出路径
    if not output_path:
        base_name, ext = os.path.splitext(file_path)
        output_path = f"{base_name}_modified{ext}"

    try:
        doc = Document(file_path)
    except Exception as e:
        return response_handler({"status": "error", "message": f"文档加载失败: {str(e)}"})

    # 解析替换指令
    replace_type = "contains"  # 默认替换类型
    search_pattern = ""
    replacement = ""
    flags = re.IGNORECASE  # 默认不区分大小写

    if ':' in replace:
        replace_type, rest = replace.split(':', 1)
        replace_type = replace_type.lower()
        if '=' in rest:
            search_pattern, replacement = rest.split('=', 1)
        else:
            return response_handler({"status": "error", "message": "无效的替换格式"})
    else:
        search_pattern, replacement = replace.split('=', 1)

    # 构建替换模式
    try:
        if replace_type == "regex":
            pattern = re.compile(search_pattern, flags)
        elif replace_type == "keyword":
            pattern = re.compile(rf'\b{re.escape(search_pattern)}\b', flags)
        elif replace_type == "contains":
            pattern = re.compile(re.escape(search_pattern), flags)
        else:
            return response_handler({"status": "error", "message": "不支持的替换类型"})
    except re.error as e:
        return response_handler({"status": "error", "message": f"正则表达式错误: {str(e)}"})

    replace_count = 0

    def process_run(run):
        """处理单个 run 对象，保留样式"""
        nonlocal replace_count
        original = run.text
        if not original:
            return

        # 替换文本
        new_text, count = pattern.subn(replacement, original)
        if count > 0:
            # 保留样式
            new_run = run._element.addprevious(run._element._new())
            new_run.text = new_text
            new_run.rPr = run._element.rPr.clone()

            # 更新 run 的文本
            run.text = ""
            replace_count += count

    def process_paragraph(para):
        """处理段落中的所有 run"""
        for run in list(para.runs):
            process_run(run)

    # 处理正文段落
    for para in doc.paragraphs:
        process_paragraph(para)

    # 处理表格
    def process_table(table):
        for row in table.rows:
            for cell in row.cells:
                # 处理单元格段落
                for para in cell.paragraphs:
                    process_paragraph(para)
                # 处理嵌套表格
                if cell.tables:
                    for nested_table in cell.tables:
                        process_table(nested_table)

    for table in doc.tables:
        process_table(table)

    # 处理页眉页脚
    for section in doc.sections:
        for header in [section.header, section.first_page_header]:
            for para in header.paragraphs:
                process_paragraph(para)
        for footer in [section.footer, section.first_page_footer]:
            for para in footer.paragraphs:
                process_paragraph(para)

    # 保存文档
    try:
        doc.save(output_path)
    except PermissionError:
        return response_handler({"status": "error", "message": "文件写入权限被拒绝"})
    except Exception as e:
        return response_handler({"status": "error", "message": f"文件保存失败: {str(e)}"})

    return response_handler({
        "status": "success",
        "data": {
            "replace_count": replace_count,
            "original_file": file_path,
            "new_file": output_path,
            "replacement_type": replace_type
        }
    })


def count_words(text, is_chinese):
    """统计文本中的字数"""
    if is_chinese:
        return len(text)
    else:
        return len(text.split())


def extract_keywords(text, is_chinese=False, top_n=10):
    """
    提取关键词,统计词频，支持中文和英文，并返回前 N 个高频词
    :param text: 文本内容
    :param is_chinese: 是否中文文档
    :param top_n: 返回的高频词数量
    :return: 关键词及其频率
    """

    if is_chinese:
        words = jieba.lcut(text)
        # 过滤标点符号和空格
        words = [word for word in words if word.strip() and not re.match(r'^[\u4e00-\u9fa5]+$', word)]
    else:
        words = re.findall(r'\b\w+\b', text.lower())

    word_counts = Counter(words)
    return word_counts.most_common(top_n)


@mcp.tool()
def extract_document_info(file_path, is_chinese=False, top_n=10, extract_content=None):
    """
    提取 Word 文档信息，包括段落、表格、图片、标题、关键词、字数
    :param file_path:
    :param is_chinese: 是否中文文档
    :param top_n: 返回的关键词数量
    :param extract_content: 要提取内容，默认全部内容
    """
    doc = Document(file_path)

    if extract_content is None:
        extract_content = ['paragraphs', 'tables', 'images', 'headings', 'text', 'keywords']

    paragraphs = []
    tables = []
    images = []
    headings = []
    text = ""
    keywords = []
    word_count = 0

    if 'paragraphs' in extract_content:
        for para in doc.paragraphs:
            para_text = para.text.strip()
            if para_text:
                paragraphs.append(para_text)
                text += para_text + " "

    if 'tables' in extract_content:
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
                text += " ".join(row_data) + " "
            tables.append(table_data)

    if 'images' in extract_content:
        for i, rel in enumerate(doc.part.rels.values()):
            if "image" in rel.target_ref:
                images.append({
                    "index": i + 1,
                    "filename": os.path.basename(rel.target_ref),
                    "path": rel.target_ref
                })

    if 'headings' in extract_content:
        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                level = int(para.style.name.replace("Heading", "").strip())
                headings.append({
                    "text": para.text.strip(),
                    "level": level
                })
                text += para.text + " "

    if 'text' in extract_content:
        word_count = count_words(text, is_chinese)

    if 'keywords' in extract_content and text:
        keywords = extract_keywords(text, is_chinese, top_n)

    result = {
        "paragraphs": paragraphs,
        "tables": tables,
        "images": images,
        "headings": headings,
        "text": text.strip(),
        "keywords": keywords,
        "word_count": word_count
    }

    return response_handler(result)


@mcp.tool()
def save_to_json(data, output_dir="output", filename="document_info.json"):
    """
    将提取的信息保存为 JSON 格式
    :param data: 提取的文档信息
    :param output_dir: 输出目录
    :param filename: 输出文件名
    """
    try:
        os.makedirs(output_dir, exist_ok=True)
        output_path = os.path.join(output_dir, filename)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=4)
    except Exception as e:
        return response_handler({"status": "error", "message": f"保存失败: {str(e)}"})

    return response_handler({"status": "success", "message": f"结果已保存到 {output_path}"})


@mcp.tool()
def save_to_csv(data, output_dir: str = "output"):
    """
    将提取的信息保存为 CSV 格式
    :param data: 提取的文档信息
    :param output_dir: 输出目录
    """

    try:
        os.makedirs(output_dir, exist_ok=True)

        # 保存段落
        if "paragraphs" in data and data["paragraphs"]:
            paragraphs_path = os.path.join(output_dir, "paragraphs.csv")
            with open(paragraphs_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["paragraph"])
                writer.writerows([[para] for para in data["paragraphs"]])

        # 保存表格
        if "tables" in data and data["tables"]:
            for i, table in enumerate(data["tables"]):
                table_path = os.path.join(output_dir, f"table_{i + 1}.csv")
                with open(table_path, "w", newline="", encoding="utf-8-sig") as f:
                    writer = csv.writer(f)
                    writer.writerows(table)

        # 保存图片信息
        if "images" in data and data["images"]:
            images_path = os.path.join(output_dir, "images.csv")
            with open(images_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.DictWriter(f, fieldnames=["index", "filename", "path"])
                writer.writeheader()
                writer.writerows(data["images"])

        # 保存标题
        if "headings" in data and data["headings"]:
            headings_path = os.path.join(output_dir, "headings.csv")
            with open(headings_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.DictWriter(f, fieldnames=["text", "level"])
                writer.writeheader()
                writer.writerows(data["headings"])

        # 保存关键词
        if "keywords" in data and data["keywords"]:
            keywords_path = os.path.join(output_dir, "keywords.csv")
            with open(keywords_path, "w", newline="", encoding="utf-8-sig") as f:
                writer = csv.writer(f)
                writer.writerow(["word", "frequency"])
                writer.writerows(data["keywords"])

        return response_handler({"status": "success", "message": f"结果已保存到 {output_dir}"})

    except Exception as e:
        return response_handler({"status": "error", "message": f"保存失败: {str(e)}"})


@mcp.tool()
def compare_documents(doc1_path: str, doc2_path: str):
    """
    比较两个文档的内容差异
    :param doc1_path: 第一个文档路径
    :param doc2_path: 第二个文档路径
    """
    try:
        doc1 = Document(doc1_path)
        doc2 = Document(doc2_path)
    except Exception as e:
        return response_handler({"error": f"无法加载文档: {str(e)}"})

    differences = {
        "added_paragraphs": [],
        "deleted_paragraphs": [],
        "modified_paragraphs": [],
        "format_differences": []
    }

    # 提取段落文本和格式
    doc1_paragraphs = []
    doc2_paragraphs = []

    for para in doc1.paragraphs:
        doc1_paragraphs.append({
            "text": para.text.strip(),
            "alignment": para.alignment,
            "font_size": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else None
        })

    for para in doc2.paragraphs:
        doc2_paragraphs.append({
            "text": para.text.strip(),
            "alignment": para.alignment,
            "font_size": para.runs[0].font.size.pt if para.runs and para.runs[0].font.size else None
        })

    # 使用 diff 算法比较段落
    differ = difflib.Differ()
    diff_result = list(differ.compare([p["text"] for p in doc1_paragraphs], [p["text"] for p in doc2_paragraphs]))

    for line in diff_result:
        if line.startswith('+'):
            differences["added_paragraphs"].append(line[2:])
        elif line.startswith('-'):
            differences["deleted_paragraphs"].append(line[2:])

    # 比较格式差异
    for i in range(min(len(doc1_paragraphs), len(doc2_paragraphs))):
        para1 = doc1_paragraphs[i]
        para2 = doc2_paragraphs[i]
        if para1["alignment"] != para2["alignment"] or para1["font_size"] != para2["font_size"]:
            differences["format_differences"].append({
                "index": i,
                "original_alignment": para1["alignment"],
                "modified_alignment": para2["alignment"],
                "original_font_size": para1["font_size"],
                "modified_font_size": para2["font_size"]
            })

    return response_handler(differences)


@mcp.tool()
def assess_document_quality(file_path: str, is_chinese: bool = False):
    """
    评估文档的质量，包括可读性和格式一致性
    :param file_path: 文档路径
    :param is_chinese: 是否中文文档
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return response_handler({"error": f"无法加载文档: {str(e)}"})

    # 提取文本内容
    text = ""
    for para in doc.paragraphs:
        text += para.text + " "

    # 可读性评估
    sentences = re.split(r'[。！？]', text) if is_chinese else re.split(r'[.!?]', text)
    avg_sentence_length = sum(len(s) for s in sentences if s) / len([s for s in sentences if s]) if sentences else 0

    # 使用更复杂的可读性指标
    readability_score = calculate_readability(text, is_chinese)

    # 格式一致性评估
    alignments = []
    font_sizes = []
    for para in doc.paragraphs:
        alignments.append(str(para.alignment))
        for run in para.runs:
            if run.font.size:
                font_sizes.append(run.font.size.pt)

    alignment_counts = Counter(alignments)
    font_size_counts = Counter(font_sizes)

    quality_report = {
        "readability": {
            "average_sentence_length": avg_sentence_length,
            "readability_score": readability_score
        },
        "consistency": {
            "paragraph_alignment": alignment_counts,
            "font_size_consistency": font_size_counts
        }
    }

    return response_handler(quality_report)


def calculate_readability(text: str, is_chinese: bool = False) -> float:
    """
    计算文本的可读性评分
    :param text: 输入文本
    :param is_chinese: 是否中文文档
    :return: 可读性评分（0-100，分数越高越易读）
    """
    if not text:
        return 0.0

    if is_chinese:
        # 中文可读性评估（示例：基于词汇复杂度和句子长度）
        import jieba
        sentences = re.split(r'[。！？]', text)
        sentences = [s for s in sentences if s.strip()]
        if not sentences:
            return 0.0

        avg_sentence_length = sum(len(s) for s in sentences) / len(sentences)
        words = [word for sentence in sentences for word in jieba.lcut(sentence)]
        unique_words = set(words)
        lexical_diversity = len(unique_words) / len(words) if words else 0

        # 示例评分公式：结合句子长度和词汇多样性
        readability_score = (100 - avg_sentence_length) * 0.5 + lexical_diversity * 50
        return min(max(readability_score, 0), 100)

    else:
        # 英文可读性评估（Flesch-Kincaid 可读性评分）
        sentences = re.split(r'[.!?]', text)
        sentences = [s for s in sentences if s.strip()]
        if not sentences:
            return 0.0

        words = re.findall(r'\b\w+\b', text.lower())
        syllables = sum(count_syllables(word) for word in words)

        avg_sentence_length = len(words) / len(sentences)
        avg_syllables_per_word = syllables / len(words) if words else 0

        # Flesch-Kincaid 公式
        readability_score = 206.835 - 1.015 * avg_sentence_length - 84.6 * avg_syllables_per_word
        return min(max(readability_score, 0), 100)


def count_syllables(word: str) -> int:
    """
    简单的音节数计数器（适用于英文）
    :param word: 单词
    :return: 音节数
    """
    if not word:
        return 0

    word = word.lower()
    vowels = "aeiouy"
    count = 0
    prev_is_vowel = False

    for char in word:
        is_vowel = char in vowels
        if is_vowel and not prev_is_vowel:
            count += 1
        prev_is_vowel = is_vowel

    # 去掉结尾的 "e" 音节
    if word.endswith("e"):
        count -= 1

    return max(count, 1)


def response_handler(response):
    return json.dumps(response, ensure_ascii=False, allow_nan=False, indent=None, separators=(',', ':'))


# 中文字符不被转义 ensure_ascii=False
# 格式紧凑 indent=None , separators=(",", ":")
# 严格性 allow_nan=False

if __name__ == "__main__":
    print("word_mcp is running...")
    mcp.run()